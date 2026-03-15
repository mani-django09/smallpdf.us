#!/usr/bin/env python3
"""
High-Fidelity PDF to Word Converter
Uses pdf2docx as primary engine (best layout preservation),
falls back to pdfplumber for table-heavy documents if needed.
"""

import sys
import os
import traceback


def is_scanned_pdf(pdf_path):
    """
    Returns True if the PDF has no extractable text (likely a scanned image PDF).
    Uses PyMuPDF (fitz) if available, otherwise assumes not scanned.
    """
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(pdf_path)
        for page in doc:
            if page.get_text().strip():
                doc.close()
                return False
        doc.close()
        return True
    except ImportError:
        # PyMuPDF not available — assume not scanned, let pdf2docx try
        return False
    except Exception:
        return False


def convert_with_pdf2docx(pdf_path, docx_path):
    """
    Primary conversion using pdf2docx.
    Preserves: fonts, images, columns, tables, text positioning, colors.
    """
    from pdf2docx import Converter

    cv = Converter(pdf_path)
    cv.convert(
        docx_path,
        start=0,
        end=None,
        zero_based_index=False,
        multi_processing=False,
        cpu_count=1,
        # Improved accuracy settings
        debug=False,
        min_section_height=20,
    )
    cv.close()
    return True


def convert_with_pdfplumber_fallback(pdf_path, docx_path):
    """
    Fallback for table-heavy PDFs (payslips, invoices, reports).
    Uses pdfplumber + python-docx with precise formatting.
    """
    import pdfplumber
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    def set_cell_shading(cell, color="D9E2F3"):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_table_borders(table, color="4472C4", size="4"):
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            tbl.insert(0, tblPr)
        for child in list(tblPr):
            if 'tblBorders' in child.tag:
                tblPr.remove(child)
        tblBorders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="{size}" w:color="{color}"/>'
            f'<w:left w:val="single" w:sz="{size}" w:color="{color}"/>'
            f'<w:bottom w:val="single" w:sz="{size}" w:color="{color}"/>'
            f'<w:right w:val="single" w:sz="{size}" w:color="{color}"/>'
            f'<w:insideH w:val="single" w:sz="{size}" w:color="{color}"/>'
            f'<w:insideV w:val="single" w:sz="{size}" w:color="{color}"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(tblBorders)

    def is_amount(text):
        if not text:
            return False
        text = str(text).strip()
        cleaned = text.replace(',', '').replace('.', '').replace('-', '').replace('/', '').replace(' ', '')
        return cleaned.isdigit() or text in ['0.00', '0', '0.0']

    doc = Document()

    # Set A4 page with narrow margins for dense documents
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            if page_num > 0:
                doc.add_page_break()

            page_text = page.extract_text() or ''
            lines = page_text.split('\n')

            # ── Extract text blocks that are NOT inside tables ──
            table_bboxes = []
            try:
                tables_raw = page.find_tables({
                    "vertical_strategy": "lines",
                    "horizontal_strategy": "lines",
                    "snap_tolerance": 3,
                })
                table_bboxes = [t.bbox for t in tables_raw] if tables_raw else []
            except Exception:
                pass

            def in_any_table(word):
                x0, top, x1, bottom = word['x0'], word['top'], word['x1'], word['bottom']
                for bbox in table_bboxes:
                    if x0 >= bbox[0] - 5 and top >= bbox[1] - 5 and x1 <= bbox[2] + 5 and bottom <= bbox[3] + 5:
                        return True
                return False

            if table_bboxes:
                non_table_words = [w for w in page.extract_words() if not in_any_table(w)]
                if non_table_words:
                    line_groups = {}
                    for word in non_table_words:
                        y_key = round(word['top'] / 6) * 6
                        if y_key not in line_groups:
                            line_groups[y_key] = []
                        line_groups[y_key].append(word)

                    for y_key in sorted(line_groups.keys()):
                        line_words = sorted(line_groups[y_key], key=lambda w: w['x0'])
                        line_text = ' '.join(w['text'] for w in line_words).strip()
                        if not line_text:
                            continue

                        para = doc.add_paragraph()
                        run = para.add_run(line_text)
                        run.font.name = 'Calibri'
                        run.font.size = Pt(10)

                        if len(line_text) < 80 and line_text.isupper():
                            run.bold = True
                            run.font.size = Pt(11)
                        elif len(line_text) < 60 and any(
                            kw in line_text.lower() for kw in
                            ['payslip', 'invoice', 'report', 'statement', 'summary', 'ltd', 'pvt', 'inc', 'private', 'limited']
                        ):
                            run.bold = True
                            run.font.size = Pt(12)
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                for line in lines:
                    line = line.strip()
                    if not line:
                        doc.add_paragraph()
                        continue
                    para = doc.add_paragraph()
                    run = para.add_run(line)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)

            # ── Process Tables ──
            table_extraction_settings = [
                {"vertical_strategy": "lines", "horizontal_strategy": "lines", "snap_tolerance": 3},
                {"vertical_strategy": "lines_strict", "horizontal_strategy": "lines_strict"},
                {"vertical_strategy": "text", "horizontal_strategy": "text", "snap_tolerance": 5},
            ]

            tables = []
            for settings in table_extraction_settings:
                try:
                    extracted = page.extract_tables(settings)
                    if extracted and any(t and len(t) > 1 for t in extracted):
                        tables = extracted
                        break
                except Exception:
                    continue

            for table_data in tables:
                if not table_data:
                    continue
                table_data = [
                    row for row in table_data
                    if row and any(cell and str(cell).strip() for cell in row)
                ]
                if not table_data:
                    continue

                num_rows = len(table_data)
                num_cols = max((len(row) for row in table_data if row), default=0)
                if num_cols == 0:
                    continue

                word_table = doc.add_table(rows=num_rows, cols=num_cols)
                word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                word_table.style = 'Table Grid'
                set_table_borders(word_table)

                for i, row in enumerate(table_data):
                    if not row:
                        continue
                    for j in range(num_cols):
                        cell_value = row[j] if j < len(row) else ''
                        cell = word_table.cell(i, j)
                        text = str(cell_value).strip() if cell_value else ''
                        cell.text = text

                        para = cell.paragraphs[0]
                        para.paragraph_format.space_before = Pt(2)
                        para.paragraph_format.space_after = Pt(2)

                        is_first_row = (i == 0)
                        is_numeric = is_amount(text)
                        font_size = 8 if num_cols > 6 else 9

                        for run in para.runs:
                            run.font.size = Pt(font_size)
                            run.font.name = 'Calibri'
                            if is_first_row:
                                run.bold = True
                            if is_numeric:
                                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                        if is_first_row:
                            set_cell_shading(cell, "BDD7EE")
                        elif text and not is_numeric and len(text) < 50 and text.istitle():
                            set_cell_shading(cell, "DDEBF7")

                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(4)

    doc.save(docx_path)
    return True


def convert_pdf_to_word(pdf_path, docx_path):
    """
    Main conversion function.
    Strategy:
      1. Detect scanned PDFs and warn early
      2. Try pdf2docx — best layout fidelity
      3. If pdf2docx fails, fall back to pdfplumber
    """
    # ── Scanned PDF detection ─────────────────────────────────────────────
    if is_scanned_pdf(pdf_path):
        print("WARNING:scanned_pdf", file=sys.stderr)
        # Still attempt conversion — pdf2docx may handle it partially

    # ── Strategy 1: pdf2docx ──────────────────────────────────────────────
    try:
        convert_with_pdf2docx(pdf_path, docx_path)
        if os.path.exists(docx_path) and os.path.getsize(docx_path) > 1000:
            print("SUCCESS:pdf2docx")
            return True
    except ImportError:
        print("WARNING: pdf2docx not installed, using fallback converter", file=sys.stderr)
    except Exception as e:
        print(f"WARNING: pdf2docx failed ({e}), trying fallback...", file=sys.stderr)

    # ── Strategy 2: pdfplumber fallback ──────────────────────────────────
    try:
        convert_with_pdfplumber_fallback(pdf_path, docx_path)
        if os.path.exists(docx_path) and os.path.getsize(docx_path) > 500:
            print("SUCCESS:pdfplumber")
            return True
    except Exception as e:
        print(f"ERROR: Fallback also failed: {e}", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)

    print("ERROR: All conversion strategies failed")
    return False


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python pdf_to_word_converter.py <input.pdf> <output.docx>")
        sys.exit(1)

    success = convert_pdf_to_word(sys.argv[1], sys.argv[2])
    sys.exit(0 if success else 1)